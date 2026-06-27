"""JobGraph - 私密求职图谱

聚焦场景: 求职 - 帮你找到靠谱的工作，避开坑人的公司
完全免费，所有功能均可使用
"""

import os
# 禁用 PyTorch 类检查警告（Streamlit 兼容性问题）
os.environ["TORCH_DISABLE_CUSTOM_CLASS_CHECK"] = "1"

import sys
from pathlib import Path
from datetime import datetime

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import streamlit as st
from loguru import logger

from src.jobgraph.graph_manager import job_manager
from src.jobgraph.models import CompanySize, UserProfile
from src.jobgraph.user.manager import user_manager
from src.jobgraph.resume import resume_parser, resume_extractor, privacy_filter, resume_optimizer
from src.jobgraph.resume.extractor import SKILL_STANDARD_MAP
from src.jobgraph.matching import job_matcher
from src.jobgraph.config_manager import config_manager
from src.jobgraph.user_data_manager import user_data_manager
from src.jobgraph.job_expectation_parser import job_expectation_parser
from src.jobgraph.notify import subscription_manager, notifier


# ============================================================
# Page Config
# ============================================================

st.set_page_config(
    page_title="JobGraph - 私密求职图谱",
    page_icon="🔒",
    layout="wide",
)

st.title("🔒 JobGraph - 私密求职图谱")
st.markdown("**你的求职，你做主** — 完全免费，数据仅存在本地")


# ============================================================
# Sidebar
# ============================================================

# 页面列表
pages = [
    "🏠 首页",
    "📄 简历管理",
    "🎯 智能匹配",
    "🔍 岗位搜索",
    "🔔 订阅提醒",
    "🏢 公司画像",
    "⚠️ 避坑指南",
    "📊 薪资行情",
    "✏️ 贡献数据",
    "🔄 数据同步",
    "👤 用户中心",
    "⚙️ LLM 配置",
]

# 获取当前页面
current_page = st.session_state.get("page", pages[0])

with st.sidebar:
    # 导航标题
    st.markdown("### 🔍 功能导航")
    st.divider()
    
    # 使用按钮实现导航，选中项有颜色差异
    for p in pages:
        if p == current_page:
            # 选中状态：使用 primary 按钮样式
            st.button(
                f"▶ {p}",
                key=f"nav_{p}",
                use_container_width=True,
                type="primary",
            )
        else:
            # 未选中状态：使用普通按钮
            if st.button(
                p,
                key=f"nav_{p}",
                use_container_width=True,
            ):
                st.session_state["page"] = p
                st.rerun()
    
    st.divider()
    
    # 免费标识
    st.success("🎁 **完全免费**")
    
    # Privacy badge
    st.info("🔒 **私密模式**\n\n数据仅存在本地")
    
    st.divider()
    
    # User info
    user_stats = user_manager.get_user_stats()
    st.markdown(f"👤 **{user_stats['nickname']}**")
    st.caption(f"等级: Lv.{user_stats['level']} | 积分: {user_stats['points']}")
    
    # Statistics
    st.markdown("### 📊 数据统计")
    try:
        stats = job_manager.get_stats()
        st.metric("🏢 公司", stats.get("companies", 0))
        st.metric("💼 岗位", stats.get("jobs", 0))
        st.metric("💬 评价", stats.get("reviews", 0))
    except Exception:
        pass

# 设置当前页面
page = current_page


# ============================================================
# Homepage
# ============================================================

if page == "🏠 首页":
    st.header("欢迎使用 JobGraph")
    
    # 隐私保护
    st.success("""
    🔒 **隐私保护** — 你的求职数据只存在你自己的电脑上
    - ✅ 无需注册，不留痕迹
    - ✅ 本地运行，数据不上传
    - ✅ 在职看机会，不会被发现
    """)
    
    # 免费标识
    st.info("""
    🎁 **完全免费** — 所有功能均可免费使用
    - ✅ 无搜索次数限制
    - ✅ 无功能限制
    - ✅ 无广告
    """)
    
    st.markdown("""
    **JobGraph** 是一款基于知识图谱的**私密求职工具**，帮你：
    - 🔍 **精准匹配** - 根据技能、薪资、地点匹配合适岗位
    - 🏢 **了解公司** - 真实员工评价，全面了解公司情况
    - ⚠️ **避坑预警** - 识别欠薪、PUA、996等坑点特征
    - 📊 **薪资分析** - 市场行情，合理定价
    """)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.info("### 🔍 岗位搜索\n聚合多平台岗位信息，智能筛选匹配")
    
    with col2:
        st.warning("### 🏢 公司画像\n真实员工评价，全面了解公司情况")
    
    with col3:
        st.error("### ⚠️ 避坑指南\n识别坑点特征，远离黑心公司")


# ============================================================
# Resume Upload
# ============================================================

elif page == "📄 简历管理":
    st.header("📄 简历管理")
    
    # 🔒 隐私保护核心提示
    st.success("""
    🔒 **简历 100% 本地处理，绝不上传服务器**
    
    - 💾 简历文件仅保存在你的电脑
    - 🔐 解析过程在本地完成，内容不会传出
    - 👤 匹配只用技能、经验，不用姓名手机号
    - 🛡️ 即使包含隐私信息，也绝不会被传出去
    """)
    
    # LLM 配置状态提示
    if not config_manager.is_llm_configured():
        st.warning("""
        ⚠️ **AI 智能解析未启用**
        
        当前使用规则模式解析简历，精度有限。
        配置 LLM 后可大幅提升解析精度。
        """)
        if st.button("⚙️ 前往配置 LLM"):
            st.session_state["page"] = "⚙️ LLM 配置"
            st.rerun()
    
    st.divider()
    
    # 获取用户 ID（用于持久化存储）
    import hashlib
    user_id = hashlib.md5(
        f"{user_manager.device_id}_resume".encode()
    ).hexdigest()[:16]
    
    # 从本地文件加载已保存的简历信息
    saved_profile = user_data_manager.load_resume_profile(user_id)
    
    if saved_profile:
        st.success("✅ 已有简历信息（本地持久化存储）")
        
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            st.info(f"当前职位: {saved_profile.get('current_title', '未识别')} | 工作年限: {saved_profile.get('experience_years', 0)} 年 | 技能: {len(saved_profile.get('skills', []))} 项")
        with col2:
            if st.button("🔄 重新上传简历"):
                user_data_manager.delete_resume_profile(user_id)
                if "resume_profile" in st.session_state:
                    del st.session_state["resume_profile"]
                st.rerun()
        with col3:
            saved_time = saved_profile.get("saved_at", "")
            if saved_time:
                st.caption(f"保存时间: {saved_time[:16]}")
        
        st.divider()
    
    # 上传区域（只有没有保存的简历时才显示）
    if not saved_profile:
        uploaded_file = st.file_uploader(
            "选择简历文件",
            type=["pdf", "docx"],
            help="支持 PDF、DOCX 格式",
        )
    else:
        uploaded_file = None
    
    # 如果有上传文件，解析它；否则使用已保存的
    profile = None
    
    if uploaded_file:
        st.success(f"✅ 已选择文件: {uploaded_file.name}")
        
        # 保存原始简历文件
        file_data = uploaded_file.getvalue()
        file_path = user_data_manager.save_resume_file(user_id, file_data, uploaded_file.name)
        if file_path:
            st.caption(f"📁 原始文件已保存: {uploaded_file.name} ({len(file_data)} bytes)")
        
        # 解析简历
        with st.spinner("正在解析简历..."):
            try:
                # 解析文件
                text = resume_parser.parse_uploaded_file(uploaded_file)
                
                # 隐私过滤
                filtered_text = privacy_filter.filter(text)
                
                # 扫描隐私信息
                privacy_scan = privacy_filter.scan(text)
                if privacy_scan:
                    st.warning(f"⚠️ 检测到隐私信息已自动过滤: {', '.join(privacy_scan.keys())}")
                
                # 提取信息
                profile = resume_extractor.extract(filtered_text)
                
                # 保存到本地文件（持久化）
                resume_data = {
                    "current_title": profile.current_title,
                    "experience_years": profile.experience_years,
                    "education": profile.education,
                    "skills": profile.skills,
                    "certifications": profile.certifications,
                    "work_history": profile.work_history,
                    "projects": profile.projects,
                    "original_filename": uploaded_file.name,
                    "file_format": Path(uploaded_file.name).suffix.lower(),
                    "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                }
                user_data_manager.save_resume_profile(user_id, resume_data)
                
                # 同时保存到 session_state
                st.session_state["resume_profile"] = resume_data
                
                st.success("✅ 简历解析完成！已保存到本地，刷新页面无需重新上传")
                
            except Exception as e:
                st.error(f"❌ 简历解析失败: {e}")
                profile = None
    elif saved_profile:
        # 使用已保存的简历信息
        profile = type('Profile', (), saved_profile)()
        # 显示提示
        st.info("📋 使用已保存的简历信息")
    
    if profile:
        st.divider()
        
        # 重新加载最新的简历数据（可能被优化建议更新过）
        saved_profile = user_data_manager.load_resume_profile(user_id)
        if saved_profile:
            profile = type('Profile', (), saved_profile)()
        
        original_filename = saved_profile.get("original_filename") if saved_profile else None
        
        if original_filename:
            st.subheader("📄 原始简历文件")
            
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                st.write(f"**文件名**: {original_filename}")
                st.write(f"**保存时间**: {saved_profile.get('saved_at', '未知')}")
            
            with col2:
                # 生成更新后的文件（保持原始格式）
                try:
                    # 获取当前简历数据（可能已更新）
                    current_profile = st.session_state.get("resume_profile", saved_profile or {})
                    current_skills = current_profile.get("skills", [])
                    file_ext = Path(original_filename).suffix.lower()
                    
                    if file_ext == ".docx":
                        # 修改原始 DOCX 文件中的技能部分
                        file_path = user_data_manager.get_resume_file_path(user_id, original_filename)
                        if file_path and file_path.exists():
                            from docx import Document as DocxDocument
                            doc = DocxDocument(str(file_path))
                            
                            # 查找并更新技能相关段落
                            skills_text = ", ".join(current_skills)
                            skill_keywords = ["技能", "技术栈", "熟悉", "精通", "掌握"]
                            
                            updated = False
                            skill_section_found = False
                            
                            for i, para in enumerate(doc.paragraphs):
                                text = para.text.strip()
                                
                                # 查找技能标题
                                if any(kw == text or f"**{kw}**" in text or f"# {kw}" in text for kw in ["技能", "技能列表", "技术栈", "Skills"]):
                                    skill_section_found = True
                                    continue
                                
                                # 如果找到了技能标题，更新下一个非空段落
                                if skill_section_found and text and len(text) > 2:
                                    # 检查是否是技能内容（包含逗号或技术词汇）
                                    if "," in text or "，" in text or any(sk.lower() in text.lower() for sk in current_skills[:3]):
                                        para.clear()
                                        para.add_run(skills_text)
                                        updated = True
                                        break
                                
                                # 如果遇到新的标题，重置
                                if text.startswith('#') or (len(text) < 20 and text.isupper()):
                                    skill_section_found = False
                            
                            # 如果没找到，在末尾添加
                            if not updated:
                                doc.add_paragraph()
                                doc.add_heading('技能', level=1)
                                doc.add_paragraph(skills_text)
                            
                            # 保存到临时文件
                            import tempfile
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                                doc.save(tmp.name)
                                with open(tmp.name, 'rb') as f:
                                    docx_data = f.read()
                                os.unlink(tmp.name)
                            
                            st.download_button(
                                label="📥 下载更新后的简历 (DOCX)",
                                data=docx_data,
                                file_name=f"{Path(original_filename).stem}_updated.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        else:
                            st.warning("原始文件不存在")
                    
                    elif file_ext == ".pdf":
                        # PDF 无法直接修改，生成 DOCX 版本
                        from docx import Document as DocxDocument
                        
                        # 读取 PDF 文本内容
                        file_path = user_data_manager.get_resume_file_path(user_id, original_filename)
                        pdf_text = ""
                        if file_path and file_path.exists():
                            try:
                                import fitz  # PyMuPDF
                                pdf_doc = fitz.open(str(file_path))
                                for page in pdf_doc:
                                    pdf_text += page.get_text()
                                pdf_doc.close()
                            except:
                                pdf_text = ""
                        
                        # 生成 DOCX（包含原始内容 + 更新后的技能）
                        doc = DocxDocument()
                        doc.add_heading('个人简历', 0)
                        
                        # 解析并添加原始内容
                        if pdf_text:
                            lines = pdf_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if not line:
                                    continue
                                # 检测标题
                                if len(line) < 20 and (line.isupper() or line.endswith('：') or line.endswith(':')):
                                    doc.add_heading(line, level=1)
                                else:
                                    doc.add_paragraph(line)
                        
                        # 更新技能部分
                        doc.add_heading('技能列表（已更新）', level=1)
                        doc.add_paragraph(', '.join(current_skills))
                        
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx_data = f.read()
                            os.unlink(tmp.name)
                        
                        st.download_button(
                            label="📥 下载更新后的简历 (DOCX)",
                            data=docx_data,
                            file_name=f"{Path(original_filename).stem}_updated.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    
                    else:
                        # 其他格式生成 DOCX
                        from docx import Document as DocxDocument
                        doc = DocxDocument()
                        doc.add_heading('个人简历', 0)
                        doc.add_heading('基本信息', level=1)
                        doc.add_paragraph(f'当前职位：{current_profile.get("current_title", "未设置")}')
                        doc.add_paragraph(f'工作年限：{current_profile.get("experience_years", 0)} 年')
                        doc.add_paragraph(f'最高学历：{current_profile.get("education", "未设置")}')
                        doc.add_heading('技能列表', level=1)
                        doc.add_paragraph(', '.join(current_skills))
                        
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx_data = f.read()
                            os.unlink(tmp.name)
                        
                        st.download_button(
                            label="📥 下载更新后的简历 (DOCX)",
                            data=docx_data,
                            file_name=f"{Path(original_filename).stem}_updated.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                except Exception as e:
                    st.warning(f"文件生成失败: {e}")
            
            with col3:
                # 预览简历
                if st.button("👁️ 预览简历", key="preview_resume_btn"):
                    st.session_state["show_resume_preview"] = not st.session_state.get("show_resume_preview", False)
            
            # 显示预览（显示更新后的内容）
            if st.session_state.get("show_resume_preview"):
                # 获取当前简历数据（可能已更新）
                current_profile = st.session_state.get("resume_profile", saved_profile or {})
                current_skills = current_profile.get("skills", [])
                
                # 显示更新后的简历内容
                st.subheader("📝 更新后的简历内容")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**当前职位**: {current_profile.get('current_title', '未识别')}")
                    st.write(f"**工作年限**: {current_profile.get('experience_years', 0)} 年")
                    st.write(f"**最高学历**: {current_profile.get('education', '未识别')}")
                
                with col2:
                    if current_skills:
                        st.write(f"**技能** ({len(current_skills)} 项):")
                        skills_text = " ".join([f"`{s}`" for s in current_skills[:15]])
                        st.markdown(skills_text)
                        if len(current_skills) > 15:
                            st.caption(f"...还有 {len(current_skills) - 15} 项技能")
                    else:
                        st.write("**技能**: 未识别")
                
                # 证书
                certifications = current_profile.get("certifications", [])
                if certifications:
                    st.write(f"**证书**: {', '.join(certifications)}")
                
                st.divider()
                
                # 原始文件下载（作为参考）
                file_path = user_data_manager.get_resume_file_path(user_id, original_filename)
                if file_path and file_path.exists():
                    with open(file_path, "rb") as f:
                        original_data = f.read()
                    st.download_button(
                        label=f"📥 下载原始文件 ({original_filename})",
                        data=original_data,
                        file_name=original_filename,
                        mime="application/octet-stream",
                    )
            
            st.divider()
        
        # 预览提取结果
        st.subheader("📋 提取结果预览")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write(f"**当前职位**: {profile.current_title or '未识别'}")
            st.write(f"**工作年限**: {profile.experience_years} 年")
            st.write(f"**最高学历**: {profile.education or '未识别'}")
        
        with col2:
            skills = profile.skills or []
            if skills:
                st.write(f"**技能** ({len(skills)} 项):")
                # 显示技能标签
                skills_text = " ".join([f"`{s}`" for s in skills[:15]])
                st.markdown(skills_text)
                if len(skills) > 15:
                    st.caption(f"...还有 {len(skills) - 15} 项技能")
            else:
                st.write("**技能**: 未识别")
        
        # 证书
        if profile.certifications:
            st.write(f"**证书**: {', '.join(profile.certifications)}")
        
        st.divider()
        
        # 用户确认/修改
        st.subheader("✏️ 确认/修改信息")
        
        with st.form("confirm_profile"):
            col1, col2 = st.columns(2)
            
            with col1:
                current_title = st.text_input(
                    "当前职位",
                    value=profile.current_title or "",
                    placeholder="如：后端工程师"
                )
                experience_years = st.number_input(
                    "工作年限",
                    min_value=0,
                    max_value=50,
                    value=profile.experience_years,
                )
                education = st.selectbox(
                    "最高学历",
                    ["大专", "本科", "硕士", "博士"],
                    index=["大专", "本科", "硕士", "博士"].index(profile.education) if profile.education in ["大专", "本科", "硕士", "博士"] else 1,
                )
            
            with col2:
                skills_input = st.text_input(
                    "技能 (逗号分隔)",
                    value=", ".join(profile.skills) if profile.skills else "",
                    placeholder="Python, Java, Go"
                )
                desired_salary = st.slider("期望薪资 (K)", 0, 100, 30)
                location = st.text_input("期望地点", placeholder="北京")
                prefer_remote = st.checkbox("接受远程办公")
            
            st.info("💡 请确认以上信息是否准确，您可以在此修改")
            
            if st.form_submit_button("🎯 开始匹配", type="primary"):
                # 使用统一的 user_id（与保存简历时一致）
                import hashlib
                
                user_id = hashlib.md5(
                    f"{user_manager.device_id}_resume".encode()
                ).hexdigest()[:16]
                
                user = UserProfile(
                    id=user_id,
                    current_title=current_title,
                    experience_years=experience_years,
                    education=education,
                    skills=[SKILL_STANDARD_MAP.get(s.strip().lower(), s.strip().title()) for s in skills_input.replace("，", ",").split(",") if s.strip()],
                    desired_salary_min=desired_salary * 1000 * 0.8 if desired_salary > 0 else None,
                    desired_salary_max=desired_salary * 1000 * 1.2 if desired_salary > 0 else None,
                    desired_locations=[loc.strip() for loc in location.replace("，", ",").split(",") if loc.strip()] if location else [],
                    prefer_remote=prefer_remote,
                    source="resume",
                    device_id=user_manager.device_id,
                )
                
                # 保存用户档案
                job_manager.create_user_profile(user)
                
                # 保存当前简历信息到 session
                st.session_state["current_profile"] = {
                    "user_id": user_id,
                    "current_title": current_title,
                    "experience_years": experience_years,
                    "education": education,
                    "skills": [s.strip() for s in skills_input.split(",") if s.strip()],
                    "work_history": profile.work_history if profile else [],
                    "projects": profile.projects if profile else [],
                }
                
                # 执行匹配
                with st.spinner("正在匹配岗位..."):
                    result = job_matcher.match_by_profile(user_id, limit=10)
                
                if result.matches:
                    # 显示匹配流程统计
                    filter_stats = result.filter_stats or {}
                    if filter_stats:
                        st.info(f"""
                        📊 **匹配流程**：
                        1️⃣ 字段初筛：从 {filter_stats.get('total_jobs', '?')} 个职位中筛选出 {filter_stats.get('filtered_count', '?')} 个
                        2️⃣ 语义匹配：{'已启用' if filter_stats.get('semantic_used') else '未启用'}，最终推荐 {filter_stats.get('final_count', '?')} 个职位
                        """)
                    
                    st.success(f"✅ 为你找到 {len(result.matches)} 个匹配岗位")
                    
                    # 显示匹配结果
                    for match in result.matches:
                        score = match.get("total_score", 0)
                        risk = match.get("company_risk", "medium")
                        
                        if score >= 0.7:
                            score_color = "green"
                        elif score >= 0.5:
                            score_color = "orange"
                        else:
                            score_color = "red"
                        
                        risk_color = "green" if risk == "low" else "orange" if risk == "medium" else "red"
                        
                        with st.expander(f"{match.get('job_title', '')} @ {match.get('company_name', '')}"):
                            # 基本信息
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                salary_min = match.get('salary_min', 0) or 0
                                salary_max = match.get('salary_max', 0) or 0
                                st.write(f"**薪资**: {salary_min/1000:.0f}-{salary_max/1000:.0f}K")
                                st.write(f"**地点**: {match.get('location', '')}")
                            
                            with col2:
                                st.markdown(f"**匹配度**: :{score_color}[{score:.0%}]")
                                st.markdown(f"**公司风险**: :{risk_color}[{risk}]")
                            
                            with col3:
                                matched_skills = match.get("matched_skills", 0)
                                st.write(f"**技能匹配**: {matched_skills}项")
                                st.write(f"**经验要求**: {match.get('experience_years', '不限')}年")
                            
                            # 技能列表
                            skills = match.get("skills", [])
                            if skills:
                                st.write(f"**技能要求**: {', '.join(skills[:10])}")
                            
                            # 岗位职责和任职要求
                            description = match.get("description", "")
                            requirements = match.get("requirements", "")
                            
                            if description or requirements:
                                st.divider()
                                
                                if description:
                                    st.write("**📋 岗位职责:**")
                                    st.markdown(description)
                                
                                if requirements:
                                    st.write("**📝 任职要求:**")
                                    st.markdown(requirements)
                            
                            if risk in ["high", "blacklist"]:
                                st.warning("⚠️ 该公司存在风险，请谨慎考虑！")
                    
                    # 如果匹配结果较少，显示简历修改建议
                    if result.need_manual_input or len(result.matches) < 5:
                        st.session_state["show_optimization"] = True
                        st.session_state["match_result"] = result
                        st.session_state["user_profile_data"] = {
                            "skills": [s.strip() for s in skills_input.split(",") if s.strip()],
                            "experience_years": experience_years,
                            "education": education,
                        }
                else:
                    st.warning("暂未找到匹配的岗位")
                    st.session_state["show_optimization"] = True
                    st.session_state["match_result"] = result
                    st.session_state["user_profile_data"] = {
                        "skills": [s.strip() for s in skills_input.split(",") if s.strip()],
                        "experience_years": experience_years,
                        "education": education,
                    }
    
    # 显示简历修改建议（在 form 外部）
    if st.session_state.get("show_optimization"):
        st.divider()
        st.subheader("📝 简历优化建议")
        
        user_profile_data = st.session_state.get("user_profile_data", {})
        match_result = st.session_state.get("match_result")
        
        # 确保 user_skills 包含用户的实际技能
        user_skills = set(user_profile_data.get("skills", []))
        if not user_skills:
            # 从保存的简历加载技能
            saved = user_data_manager.load_resume_profile(user_id)
            if saved:
                user_skills = set(saved.get("skills", []))
                user_profile_data["skills"] = list(user_skills)
        
        # 检查是否有匹配结果
        has_matches = match_result and match_result.matches and len(match_result.matches) > 0
        
        if not has_matches:
            # 没有匹配结果，分析原因并提供建议
            st.warning("⚡ **匹配结果为空**")
            
            # 分析可能的原因
            st.write("**可能原因分析：**")
            
            user_locations = st.session_state.get("user_profile_data", {}).get("desired_locations", [])
            
            # 查询数据库中的热门技能
            from src.graph.neo4j_client import neo4j_client
            hot_skills = []
            try:
                result = neo4j_client.execute_query('''
                    MATCH (j:Job)
                    WHERE j.skills IS NOT NULL
                    UNWIND j.skills AS skill
                    RETURN skill, count(*) AS cnt
                    ORDER BY cnt DESC
                    LIMIT 30
                ''')
                all_hot_skills = [r["skill"] for r in result]
                
                # 计算技能重叠
                overlap = user_skills & set(all_hot_skills)
                
                if not overlap:
                    st.write("1. ❌ **技能不匹配**：您的技能与当前职位需求不匹配")
                    st.write(f"   热门技能: {', '.join(all_hot_skills[:10])}")
                else:
                    st.write(f"1. ✅ 技能匹配: {', '.join(overlap)}")
                
                # 过滤掉用户已有的技能
                hot_skills = [s for s in all_hot_skills if s not in user_skills]
            except Exception as e:
                logger.error(f"查询热门技能失败: {e}")
            
            if user_locations:
                st.write(f"2. 📍 地点限制: {', '.join(user_locations)}")
                st.write("   建议：扩大地点范围或留空")
            
            st.divider()
            
            # 自动更新简历建议
            st.write("**🔧 自动更新简历：**")
            
            # 获取推荐技能（排除用户已有的）
            hot_skills_filtered = [s for s in hot_skills[:10] if s not in user_skills] if hot_skills else []
            
            # 显示状态信息
            if not hot_skills:
                st.warning("⚠️ 数据库中暂无职位数据，请先同步数据或添加职位")
            elif not user_skills:
                st.info("💡 未获取到您的技能信息，请先在「简历管理」页面上传简历")
            elif not hot_skills_filtered:
                st.success("✅ 您的技能已涵盖当前热门技能！")
                st.write("建议放宽筛选条件或查看其他职位")
            else:
                # 有可选技能，让用户选择
                st.write("选择要添加到简历的热门技能：")
                
                # 初始化默认选中状态（只在首次）
                for i, skill in enumerate(hot_skills_filtered[:9]):
                    key = f"skill_{skill}"
                    if key not in st.session_state:
                        st.session_state[key] = (i < 3)
                
                # 用户勾选技能（不用 columns，避免状态问题）
                selected_skills = []
                for skill in hot_skills_filtered[:9]:
                    if st.checkbox(skill, key=f"skill_{skill}"):
                        selected_skills.append(skill)
                
                # 实时显示预览
                if selected_skills:
                    suggested = list(user_skills) + selected_skills
                    st.info(f"添加后技能：`{'`, `'.join(suggested[:15])}`")
                    
                    if st.button("✅ 更新简历", type="primary", key="auto_update_empty"):
                        import hashlib
                        user_id = hashlib.md5(
                            f"{user_manager.device_id}_resume".encode()
                        ).hexdigest()[:16]
                        
                        saved_profile = user_data_manager.load_resume_profile(user_id) or {}
                        
                        # 构建更新后的简历数据（保留原始内容，更新技能）
                        updated_resume = {
                            **saved_profile,
                            "skills": suggested,
                            "current_title": user_profile_data.get("current_title", saved_profile.get("current_title", "")),
                            "experience_years": user_profile_data.get("experience_years", saved_profile.get("experience_years", 0)),
                            "education": user_profile_data.get("education", saved_profile.get("education")),
                            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        }
                        
                        # 保存到文件
                        user_data_manager.save_resume_profile(user_id, updated_resume)
                        
                        # 同时更新 session_state（用于预览和下载）
                        st.session_state["resume_profile"] = updated_resume
                        
                        # 更新 UserProfile
                        user = UserProfile(
                            id=user_id,
                            current_title=updated_resume.get("current_title", ""),
                            experience_years=updated_resume.get("experience_years", 0),
                            education=updated_resume.get("education"),
                            skills=suggested,
                            source="resume",
                            device_id=user_manager.device_id,
                        )
                        job_manager.create_user_profile(user)
                        
                        st.success("✅ 简历已更新！请重新匹配")
                        st.session_state["show_optimization"] = False
                        st.rerun()
                else:
                    st.warning("请至少选择一项技能")
            
            st.divider()
            st.write("**💡 其他建议：**")
            st.markdown("""
            1. **放宽地点限制**：将期望地点留空，匹配所有城市
            2. **手动输入职位信息**：使用「智能匹配」页面手动输入
            3. **查看所有职位**：在「岗位搜索」页面浏览所有职位
            """)
            
            if st.button("📝 手动输入职位信息", key="manual_from_empty"):
                st.session_state["page"] = "🎯 智能匹配"
                st.session_state["show_optimization"] = False
                st.rerun()
        else:
            # 有匹配结果，分析技能差距
            all_job_skills = set()
            for job in match_result.matches[:3]:
                all_job_skills.update(job.get("skills", []))
            
            missing_skills = [s for s in all_job_skills if s not in user_skills]
            
            if missing_skills:
                st.warning(f"⚡ **技能差距**：匹配职位需要以下技能，您可以选择添加：")
                
                # 初始化默认选中状态
                for i, skill in enumerate(sorted(missing_skills)[:9]):
                    key = f"match_skill_{skill}"
                    if key not in st.session_state:
                        st.session_state[key] = (i < 3)
                
                # 用户勾选技能
                selected_skills = []
                for skill in sorted(missing_skills)[:9]:
                    if st.checkbox(skill, key=f"match_skill_{skill}"):
                        selected_skills.append(skill)
                
                # 实时显示预览
                if selected_skills:
                    current_skills = list(user_skills)
                    suggested_skills = current_skills + selected_skills
                    st.info(f"添加后技能：`{'`, `'.join(suggested_skills[:15])}`")
                    
                    if st.button("✅ 更新简历", type="primary", key="update_from_match"):
                        import hashlib
                        user_id = hashlib.md5(
                            f"{user_manager.device_id}_resume".encode()
                        ).hexdigest()[:16]
                        
                        saved_profile = user_data_manager.load_resume_profile(user_id) or {}
                        
                        # 构建更新后的简历数据（保留原始内容，更新技能）
                        updated_resume = {
                            **saved_profile,
                            "skills": suggested_skills,
                            "current_title": user_profile_data.get("current_title", saved_profile.get("current_title", "")),
                            "experience_years": user_profile_data.get("experience_years", saved_profile.get("experience_years", 0)),
                            "education": user_profile_data.get("education", saved_profile.get("education")),
                            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        }
                        
                        # 保存到文件
                        user_data_manager.save_resume_profile(user_id, updated_resume)
                        
                        # 同时更新 session_state
                        st.session_state["resume_profile"] = updated_resume
                        
                        # 更新 UserProfile
                        user = UserProfile(
                            id=user_id,
                            current_title=updated_resume.get("current_title", ""),
                            experience_years=updated_resume.get("experience_years", 0),
                            education=updated_resume.get("education"),
                            skills=suggested_skills,
                            source="resume",
                            device_id=user_manager.device_id,
                        )
                        job_manager.create_user_profile(user)
                        
                        st.success("✅ 简历已更新！请重新匹配")
                        st.session_state["show_optimization"] = False
                        st.rerun()
                else:
                    st.warning("请至少选择一项技能")
    
    # 显示更新后的简历（格式与更新前一致）
    if st.session_state.get("show_updated_resume"):
        st.divider()
        st.subheader("📄 更新后的简历")
        
        updated = st.session_state.get("updated_resume", {})
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write(f"**当前职位**: {updated.get('current_title', '未识别')}")
            st.write(f"**工作年限**: {updated.get('experience_years', 0)} 年")
            st.write(f"**最高学历**: {updated.get('education', '未识别')}")
        
        with col2:
            skills = updated.get("skills", [])
            if skills:
                st.write(f"**技能** ({len(skills)} 项):")
                # 使用标签格式显示技能（与更新前一致）
                skills_text = " ".join([f"`{s}`" for s in skills[:15]])
                st.markdown(skills_text)
                if len(skills) > 15:
                    st.caption(f"...还有 {len(skills) - 15} 项技能")
            else:
                st.write("**技能**: 未识别")
        
        # 证书
        certifications = updated.get("certifications", [])
        if certifications:
            st.write(f"**证书**: {', '.join(certifications)}")
        
        st.caption(f"更新时间: {updated.get('updated_at', '')}")
        
        st.divider()
        
        # 生成可下载的简历
        file_format = updated.get("file_format", ".txt")
        original_filename = updated.get("original_filename", "resume")
        base_name = Path(original_filename).stem if original_filename else "resume"
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # TXT 格式下载
            resume_text = f"""个人简历
========

基本信息
--------
当前职位：{updated.get('current_title', '未设置')}
工作年限：{updated.get('experience_years', 0)} 年
最高学历：{updated.get('education', '未设置')}

技能列表
--------
{', '.join(skills)}

更新时间：{updated.get('updated_at', '')}
"""
            st.download_button(
                label="📥 下载简历 (TXT)",
                data=resume_text,
                file_name=f"{base_name}_updated.txt",
                mime="text/plain",
            )
        
        with col2:
            # DOCX 格式下载
            if file_format == ".docx" or st.button("📥 下载简历 (DOCX)"):
                try:
                    from docx import Document as DocxDocument
                    
                    doc = DocxDocument()
                    doc.add_heading('个人简历', 0)
                    
                    # 基本信息
                    doc.add_heading('基本信息', level=1)
                    doc.add_paragraph(f'当前职位：{updated.get("current_title", "未设置")}')
                    doc.add_paragraph(f'工作年限：{updated.get("experience_years", 0)} 年')
                    doc.add_paragraph(f'最高学历：{updated.get("education", "未设置")}')
                    
                    # 技能列表
                    doc.add_heading('技能列表', level=1)
                    doc.add_paragraph(', '.join(skills))
                    
                    # 证书
                    if certifications:
                        doc.add_heading('证书', level=1)
                        doc.add_paragraph(', '.join(certifications))
                    
                    # 保存到临时文件
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx_data = f.read()
                        os.unlink(tmp.name)
                    
                    st.download_button(
                        label="📥 下载简历 (DOCX)",
                        data=docx_data,
                        file_name=f"{base_name}_updated.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                except Exception as e:
                    st.warning(f"DOCX 生成失败: {e}")
        
        with col3:
            if st.button("🔄 重新匹配", type="primary"):
                st.session_state["show_updated_resume"] = False
                st.rerun()
    
    # 手动输入按钮（在 form 外部）
    if st.session_state.get("show_manual_btn"):
        st.divider()
        st.info("💡 建议尝试手动输入职位信息进行匹配")
        if st.button("📝 手动输入职位信息", key="manual_from_empty"):
            st.session_state["page"] = "🎯 智能匹配"
            st.session_state["show_optimization"] = False
            st.rerun()
    
    else:
        # 未上传文件时显示说明
        st.info("""
        **使用说明**：
        1. 上传您的简历文件（PDF 或 DOCX 格式）
        2. 系统自动解析简历内容
        3. 确认提取的信息是否准确
        4. 点击"开始匹配"获取推荐岗位
        """)
        
        st.markdown("""
        **支持的信息提取**：
        - 🔧 技能（编程语言、框架、工具等）
        - 📅 工作年限
        - 🎓 教育背景
        - 💼 当前职位
        - 📜 专业证书
        """)

# ============================================================
# Subscription & Notifications
# ============================================================

elif page == "🔔 订阅提醒":
    st.header("🔔 订阅提醒")
    
    st.info("订阅关键词，有新职位时自动提醒")
    
    # 获取用户 ID
    import hashlib
    user_id = hashlib.md5(f"{user_manager.device_id}_sub".encode()).hexdigest()[:16]
    
    # 添加订阅
    st.subheader("📝 添加订阅")
    
    with st.form("add_subscription"):
        col1, col2 = st.columns(2)
        
        with col1:
            keywords = st.text_input(
                "关键词 (逗号分隔)",
                placeholder="Java, Python, 后端",
                help="职位标题或技能中包含这些关键词时通知您"
            )
            city = st.text_input("城市", value="武汉", placeholder="武汉")
        
        with col2:
            salary_min = st.number_input("最低薪资 (K)", 0, 100, 0, help="0 表示不限")
            notify_method = st.selectbox("通知方式", ["应用内通知", "邮件", "微信"])
        
        if st.form_submit_button("✅ 添加订阅", type="primary"):
            if keywords:
                sub = subscription_manager.add_subscription(
                    user_id=user_id,
                    keywords=[k.strip() for k in keywords.split(",") if k.strip()],
                    city=city,
                    salary_min=salary_min * 1000 if salary_min > 0 else None,
                )
                st.success(f"✅ 订阅成功！关键词: {', '.join(sub['keywords'])}")
                st.rerun()
            else:
                st.warning("请输入关键词")
    
    st.divider()
    
    # 订阅列表
    st.subheader("📋 我的订阅")
    
    subs = subscription_manager.get_user_subscriptions(user_id)
    
    if subs:
        for sub in subs:
            col1, col2, col3 = st.columns([3, 2, 1])
            
            with col1:
                st.write(f"**关键词**: {', '.join(sub['keywords'])}")
            with col2:
                st.write(f"**城市**: {sub.get('city', '不限')}")
            with col3:
                if st.button("🗑️", key=f"del_sub_{sub['id']}"):
                    subscription_manager.remove_subscription(sub['id'])
                    st.rerun()
    else:
        st.info("暂无订阅，添加订阅后有新职位时会自动提醒")
    
    st.divider()
    
    # 通知列表
    st.subheader("🔔 新职位通知")
    
    # 标记全部已读按钮
    col1, col2 = st.columns([3, 1])
    with col2:
        if st.button("全部已读"):
            notifier.mark_all_read(user_id)
            st.rerun()
    
    notifs = notifier.get_user_notifications(user_id)
    
    if notifs:
        for notif in notifs[:20]:
            read_icon = "📖" if notif["read"] else "📩"
            
            with st.expander(f"{read_icon} {notif['job_title']} @ {notif['company_name']}"):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write(f"**职位**: {notif['job_title']}")
                    st.write(f"**公司**: {notif['company_name']}")
                    st.write(f"**地点**: {notif.get('location', '未知')}")
                
                with col2:
                    salary_min = notif.get('salary_min') or 0
                    salary_max = notif.get('salary_max') or 0
                    if salary_min or salary_max:
                        st.write(f"**薪资**: {salary_min/1000:.0f}-{salary_max/1000:.0f}K")
                    
                    skills = notif.get('skills', [])
                    if skills:
                        st.write(f"**技能**: {', '.join(skills[:5])}")
                
                # 标记已读
                if not notif["read"]:
                    notifier.mark_read(notif["id"])
    else:
        st.info("暂无通知")
    
    # 订阅统计
    st.divider()
    sub_stats = subscription_manager.get_stats()
    notif_stats = notifier.get_stats(user_id)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("活跃订阅", sub_stats["active_subscriptions"])
    with col2:
        st.metric("总通知", notif_stats["total"])
    with col3:
        st.metric("未读通知", notif_stats["unread"])


elif page == "🔍 岗位搜索":
    st.header("🔍 岗位搜索")
    
    # Search filters
    col1, col2, col3 = st.columns(3)
    
    with col1:
        query = st.text_input("搜索关键词", placeholder="职位名称/公司")
    
    with col2:
        location = st.selectbox("工作地点", ["不限", "北京", "上海", "广州", "深圳", "杭州", "成都"])
    
    with col3:
        salary_min = st.slider("最低薪资 (K)", 0, 100, 0)
    
    # Search
    if st.button("搜索", type="primary") or query:
        try:
            jobs = job_manager.search_jobs(
                query=query if query else None,
                location=location if location != "不限" else None,
                salary_min=salary_min * 1000 if salary_min > 0 else None,
                limit=50
            )
            
            if jobs:
                st.subheader(f"找到 {len(jobs)} 个岗位")
                
                for job in jobs:
                    j = job.get("j", {})
                    risk = job.get("company_risk", "medium")
                    risk_color = "green" if risk == "low" else "orange" if risk == "medium" else "red"
                    risk_text = {"low": "低风险", "medium": "中风险", "high": "高风险", "blacklist": "黑名单"}.get(risk, "未知")
                    
                    with st.expander(f"{j.get('title', '')} @ {j.get('company_name', '')}"):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.write(f"**薪资**: {j.get('salary_min', 0)/1000:.0f}-{j.get('salary_max', 0)/1000:.0f}K × {j.get('salary_months', 12)}薪")
                            st.write(f"**地点**: {j.get('location', '')}")
                            st.write(f"**经验**: {j.get('experience_years', 0)}年+")
                            st.write(f"**学历**: {j.get('education', '不限')}")
                        
                        with col2:
                            st.markdown(f"**公司风险**: :{risk_color}[{risk_text}]")
                            skills = j.get("skills", [])
                            if skills:
                                st.write(f"**技能要求**: {', '.join(skills)}")
                        
                        benefits = j.get("benefits", [])
                        if benefits:
                            st.write(f"**福利待遇**: {', '.join(benefits)}")
                        
                        if risk in ["high", "blacklist"]:
                            st.warning("⚠️ 该公司存在风险，请谨慎考虑！")
            else:
                st.info("未找到匹配的岗位")
        except Exception as e:
            st.error(f"搜索失败: {e}")


# ============================================================
# Company Profile
# ============================================================

elif page == "🏢 公司画像":
    st.header("🏢 公司画像")
    
    company_query = st.text_input("搜索公司", placeholder="输入公司名称")
    
    if company_query:
        companies = job_manager.search_companies(company_query)
        
        if companies:
            for comp in companies:
                c = comp.get("c", {})
                risk = c.get("risk_level", "medium")
                risk_icon = "🟢" if risk == "low" else "🟡" if risk == "medium" else "🔴"
                
                st.subheader(f"{risk_icon} {c.get('name', '')}")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.write(f"**行业**: {c.get('industry', '')}")
                    st.write(f"**规模**: {c.get('size', '')}")
                    st.write(f"**成立**: {c.get('founded', '')}")
                    st.write(f"**总部**: {c.get('headquarters', '')}")
                
                with col2:
                    st.write(f"**融资**: {c.get('funding_stage', '')}")
                    st.write(f"**上市**: {'是' if c.get('is_listed') else '否'}")
                    st.write(f"**平均薪资**: ¥{c.get('avg_salary', 0):,.0f}/月")
                    st.write(f"**评分**: {'⭐' * int(c.get('avg_rating', 0))} ({c.get('avg_rating', 0):.1f})")
                
                with col3:
                    st.markdown(f"**风险等级**: {risk.upper()}")
                    st.write(f"**风险分数**: {c.get('risk_score', 0):.2f}")
                    
                    tags = c.get("tags", [])
                    if tags:
                        st.write(f"**标签**: {', '.join(tags)}")
                
                # Reviews
                reviews = job_manager.get_company_reviews(c.get("id", ""))
                if reviews:
                    st.subheader("💬 员工评价")
                    for rev in reviews[:5]:
                        r = rev.get("r", {})
                        st.write(f"**{r.get('title', '')}** - {r.get('reviewer_title', '')}")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.success(f"✅ {r.get('pros', '')}")
                        with col2:
                            st.error(f"❌ {r.get('cons', '')}")
                
                # Pitfalls
                pitfalls = job_manager.get_company_pitfalls(c.get("id", ""))
                if pitfalls:
                    st.subheader("⚠️ 坑点预警")
                    for pit in pitfalls:
                        p = pit.get("p", {})
                        severity = p.get("severity", 3)
                        st.error(f"**{p.get('pitfall_type', '')}** ({'🔴' * severity})")
                
                st.divider()
        else:
            st.info("未找到匹配的公司")


# ============================================================
# Pitfall Guide
# ============================================================

elif page == "⚠️ 避坑指南":
    st.header("⚠️ 避坑指南")
    
    st.subheader("常见坑点类型")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.error("### 💰 欠薪\n拖欠工资，不按时发放")
    
    with col2:
        st.error("### 🎭 PUA\n精神控制，贬低员工")
    
    with col3:
        st.error("### ⏰ 996\n强制加班，没有加班费")
    
    with col4:
        st.error("### 📉 裁员\n频繁裁员，不稳定")
    
    st.divider()
    
    st.subheader("避坑技巧")
    
    st.info("""
    1. **查看员工评价**: 脉脉、Glassdoor 等平台的真实评价
    2. **查询企业信息**: 天眼查、企查查了解公司背景
    3. **注意面试信号**: 面试中是否尊重候选人
    4. **问清楚薪资结构**: 底薪、绩效、年终奖比例
    5. **了解加班情况**: 是否强制加班，有无加班费
    """)
    
    st.divider()
    
    st.subheader("黑名单公司查询")
    
    blacklist_query = st.text_input("查询公司", placeholder="输入公司名称")
    
    if blacklist_query:
        companies = job_manager.search_companies(blacklist_query)
        
        for comp in companies:
            c = comp.get("c", {})
            risk = c.get("risk_level", "medium")
            
            if risk in ["high", "blacklist"]:
                st.error(f"⚠️ **{c.get('name', '')}** 存在高风险！")
                st.write(f"风险等级: {risk.upper()}")
                st.write(f"风险分数: {c.get('risk_score', 0):.2f}")
            else:
                st.success(f"✅ **{c.get('name', '')}** 暂未发现明显风险")


# ============================================================
# Salary Analysis
# ============================================================

elif page == "📊 薪资行情":
    st.header("📊 薪资行情")
    
    st.subheader("薪资查询")
    
    job_title = st.text_input("输入职位名称", placeholder="如: 后端工程师")
    
    if job_title:
        salary_data = job_manager.get_job_salary_range(job_title)
        
        if salary_data and salary_data.get("sample_count", 0) > 0:
            st.subheader(f"「{job_title}」薪资分布")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("25分位", f"¥{salary_data.get('p25_min', 0)/1000:.0f}K")
            
            with col2:
                st.metric("中位数", f"¥{salary_data.get('p50_min', 0)/1000:.0f}K")
            
            with col3:
                st.metric("75分位", f"¥{salary_data.get('p75_min', 0)/1000:.0f}K")
            
            st.write(f"样本数量: {salary_data.get('sample_count', 0)}")
        else:
            st.info(f"暂无「{job_title}」的薪资数据")


# ============================================================
# Smart Matching
# ============================================================

elif page == "🎯 智能匹配":
    st.header("🎯 智能匹配")
    
    st.info("填写你的期望职位信息，系统会为你推荐匹配的岗位")
    
    # 隐私提示
    st.warning("🔒 **隐私保护**：无需填写姓名、手机号等隐私信息")
    
    st.divider()
    
    # 输入方式选择
    input_method = st.radio(
        "选择输入方式",
        ["📝 手动填写", "📋 粘贴文本", "🖼️ 上传图片"],
        horizontal=True,
    )
    
    # 初始化表单默认值
    default_job_title = ""
    default_skills = ""
    default_experience_years = 3
    default_education = "本科"
    default_desired_salary = 30
    default_location = ""
    default_prefer_remote = False
    default_summary = ""
    
    # 根据输入方式解析
    if input_method == "📋 粘贴文本":
        st.subheader("📋 粘贴职位期望文本")
        
        input_text = st.text_area(
            "粘贴文本",
            placeholder="示例：\n期望职位：后端工程师\n技能要求：Python, Java, MySQL, Redis\n工作年限：5年以上\n期望薪资：30-50K\n工作地点：北京",
            height=150,
        )
        
        if input_text:
            # 解析文本
            parsed = job_expectation_parser.parse_text(input_text)
            
            if parsed:
                st.success(f"✅ 解析成功！识别到 {len(parsed)} 个字段")
                
                # 显示解析结果
                col1, col2 = st.columns(2)
                with col1:
                    if parsed.get("job_title"):
                        st.write(f"**职位**: {parsed['job_title']}")
                    if parsed.get("experience_years"):
                        st.write(f"**年限**: {parsed['experience_years']} 年")
                    if parsed.get("education"):
                        st.write(f"**学历**: {parsed['education']}")
                
                with col2:
                    if parsed.get("skills"):
                        st.write(f"**技能**: {', '.join(parsed['skills'])}")
                    if parsed.get("location"):
                        st.write(f"**地点**: {parsed['location']}")
                    if parsed.get("salary_min"):
                        st.write(f"**薪资**: {parsed['salary_min']/1000:.0f}-{parsed.get('salary_max', 0)/1000:.0f}K")
                
                # 显示个人简介
                if parsed.get("summary"):
                    st.write(f"**个人简介**: {parsed['summary'][:100]}...")
                
                # 更新默认值
                default_job_title = parsed.get("job_title", "")
                default_skills = ", ".join(parsed.get("skills", []))
                default_experience_years = parsed.get("experience_years", 3)
                default_education = parsed.get("education", "本科")
                default_desired_salary = int(parsed.get("salary_min", 30000) / 1000)
                default_location = parsed.get("location", "")
                default_summary = parsed.get("summary", "")
            else:
                st.warning("未能解析出有效信息，请检查文本格式或手动填写")
    
    elif input_method == "🖼️ 上传图片":
        st.subheader("🖼️ 上传职位期望图片")
        
        uploaded_image = st.file_uploader(
            "选择图片",
            type=["png", "jpg", "jpeg"],
            help="支持 PNG、JPG 格式",
        )
        
        if uploaded_image:
            st.image(uploaded_image, caption="上传的图片", use_column_width=True)
            
            with st.spinner("正在识别图片文字..."):
                try:
                    # 使用 OCR 识别图片文字
                    import easyocr
                    reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
                    
                    # 读取图片
                    import numpy as np
                    from PIL import Image
                    image = Image.open(uploaded_image)
                    img_array = np.array(image)
                    
                    # OCR 识别
                    results = reader.readtext(img_array)
                    ocr_text = " ".join([r[1] for r in results])
                    
                    st.write("**识别结果**:")
                    st.code(ocr_text)
                    
                    # 解析 OCR 文本
                    parsed = job_expectation_parser.parse_text(ocr_text)
                    
                    if parsed:
                        st.success(f"✅ 解析成功！识别到 {len(parsed)} 个字段")
                        
                        # 更新默认值
                        default_job_title = parsed.get("job_title", "")
                        default_skills = ", ".join(parsed.get("skills", []))
                        default_experience_years = parsed.get("experience_years", 3)
                        default_education = parsed.get("education", "本科")
                        default_desired_salary = int(parsed.get("salary_min", 30000) / 1000)
                        default_location = parsed.get("location", "")
                        default_summary = parsed.get("summary", "")
                    else:
                        st.warning("未能解析出有效信息，请手动填写")
                        
                except Exception as e:
                    st.error(f"图片识别失败: {e}")
                    st.info("请手动填写或尝试粘贴文本")
    
    st.divider()
    
    # 匹配表单（始终显示）
    st.subheader("📋 期望职位信息")
    
    with st.form("job_expectation_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            job_title = st.text_input("期望职位", value=default_job_title, placeholder="如：后端工程师")
            skills = st.text_input("技能 (逗号分隔)", value=default_skills, placeholder="Python, Java, Go")
            experience_years = st.number_input("工作年限", 0, 30, value=default_experience_years)
            education = st.selectbox(
                "学历要求",
                ["不限", "大专", "本科", "硕士", "博士"],
                index=["不限", "大专", "本科", "硕士", "博士"].index(default_education) if default_education in ["不限", "大专", "本科", "硕士", "博士"] else 2,
            )
        
        with col2:
            desired_salary = st.slider("期望薪资 (K)", 0, 100, value=default_desired_salary)
            location = st.text_input("期望地点", value=default_location, placeholder="北京")
            prefer_remote = st.checkbox("接受远程办公", value=default_prefer_remote)
        
        # 个人简介输入
        summary = st.text_area(
            "个人简介 (可选)",
            value=default_summary,
            placeholder="简单介绍自己的工作经验、技术栈、项目经验等，用于语义匹配...",
            height=100,
            help="填写个人简介可以提高匹配精度，系统会分析简介内容与职位描述的匹配度"
        )
        
        if st.form_submit_button("🎯 开始匹配", type="primary"):
            if not job_title and not skills:
                st.warning("请至少填写期望职位或技能")
            else:
                import hashlib
                
                user_id = hashlib.md5(
                    f"{user_manager.device_id}_smart".encode()
                ).hexdigest()[:16]
                
                user = UserProfile(
                    id=user_id,
                    current_title=job_title,
                    experience_years=experience_years,
                    education=education if education != "不限" else None,
                    skills=[SKILL_STANDARD_MAP.get(s.strip().lower(), s.strip().title()) for s in skills.replace("，", ",").split(",") if s.strip()],
                    desired_salary_min=desired_salary * 1000 * 0.8 if desired_salary > 0 else None,
                    desired_salary_max=desired_salary * 1000 * 1.2 if desired_salary > 0 else None,
                    desired_locations=[loc.strip() for loc in location.replace("，", ",").split(",") if loc.strip()] if location else [],
                    prefer_remote=prefer_remote,
                    resume_text=summary if summary else None,
                    source="smart",
                    device_id=user_manager.device_id,
                )
                
                job_manager.create_user_profile(user)
                
                with st.spinner("正在匹配岗位..."):
                    result = job_matcher.match_by_profile(user_id, limit=10)
                
                if result.matches:
                    # 显示匹配流程统计
                    filter_stats = result.filter_stats or {}
                    if filter_stats:
                        st.info(f"""
                        📊 **匹配流程**：
                        1️⃣ 字段初筛：从 {filter_stats.get('total_jobs', '?')} 个职位中筛选出 {filter_stats.get('filtered_count', '?')} 个
                        2️⃣ 语义匹配：{'已启用' if filter_stats.get('semantic_used') else '未启用'}，最终推荐 {filter_stats.get('final_count', '?')} 个职位
                        """)
                    
                    st.success(f"✅ 为你找到 {len(result.matches)} 个匹配岗位")
                    
                    # 显示匹配结果
                    for match in result.matches:
                        score = match.get("total_score", 0)
                        risk = match.get("company_risk", "medium")
                        
                        if score >= 0.7:
                            score_color = "green"
                        elif score >= 0.5:
                            score_color = "orange"
                        else:
                            score_color = "red"
                        
                        risk_color = "green" if risk == "low" else "orange" if risk == "medium" else "red"
                        
                        with st.expander(f"{match.get('job_title', '')} @ {match.get('company_name', '')}"):
                            # 基本信息
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                salary_min = match.get('salary_min', 0) or 0
                                salary_max = match.get('salary_max', 0) or 0
                                st.write(f"**薪资**: {salary_min/1000:.0f}-{salary_max/1000:.0f}K")
                                st.write(f"**地点**: {match.get('location', '')}")
                            
                            with col2:
                                st.markdown(f"**匹配度**: :{score_color}[{score:.0%}]")
                                st.markdown(f"**公司风险**: :{risk_color}[{risk}]")
                            
                            with col3:
                                matched_skills = match.get("matched_skills", 0)
                                st.write(f"**技能匹配**: {matched_skills}项")
                                st.write(f"**经验要求**: {match.get('experience_years', '不限')}年")
                            
                            # 技能列表
                            skills = match.get("skills", [])
                            if skills:
                                st.write(f"**技能要求**: {', '.join(skills[:10])}")
                            
                            # 岗位职责和任职要求
                            description = match.get("description", "")
                            requirements = match.get("requirements", "")
                            
                            if description or requirements:
                                st.divider()
                                
                                if description:
                                    st.write("**📋 岗位职责:**")
                                    st.markdown(description)
                                
                                if requirements:
                                    st.write("**📝 任职要求:**")
                                    st.markdown(requirements)
                            
                            if risk in ["high", "blacklist"]:
                                st.warning("⚠️ 该公司存在风险，请谨慎考虑！")
                else:
                    st.warning("暂未找到匹配的岗位")
                    st.info("💡 建议尝试：")
                    st.markdown("""
                    - 调整期望职位关键词
                    - 减少筛选条件
                    - 扩大地点范围
                    - 降低薪资要求
                    """)

# ============================================================
# Contribution Page
# ============================================================

elif page == "✏️ 贡献数据":
    st.header("✏️ 贡献数据")
    
    st.info("帮助其他求职者了解公司真实情况，提交你的评价和坑点")
    
    # 用户贡献统计
    user_stats = user_manager.get_user_stats()
    contributions = user_stats.get("contributions", {})
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("等级", f"Lv.{user_stats['level']}")
    with col2:
        st.metric("积分", user_stats['points'])
    with col3:
        st.metric("评价", contributions.get('reviews', 0))
    with col4:
        st.metric("坑点", contributions.get('pitfalls', 0))
    
    st.divider()
    
    tab1, tab2, tab3, tab4 = st.tabs(["提交评价", "提交坑点", "提交薪资", "提交职位"])
    
    with tab1:
        st.subheader("提交员工评价")
        
        company_query = st.text_input("搜索公司", placeholder="输入公司名称", key="review_company")
        
        if company_query:
            companies = job_manager.search_companies(company_query)
            if companies:
                company_options = [c.get("c", {}).get("name", "") for c in companies]
                selected_company = st.selectbox("选择公司", company_options)
                
                company_id = None
                for c in companies:
                    if c.get("c", {}).get("name") == selected_company:
                        company_id = c.get("c", {}).get("id")
                        break
                
                if company_id:
                    with st.form("review_form"):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            overall_rating = st.slider("综合评分", 1.0, 5.0, 3.0, 0.1)
                            salary_rating = st.slider("薪资评分", 1.0, 5.0, 3.0, 0.1)
                            work_life_rating = st.slider("工作生活平衡", 1.0, 5.0, 3.0, 0.1)
                        
                        with col2:
                            management_rating = st.slider("管理评分", 1.0, 5.0, 3.0, 0.1)
                            reviewer_title = st.text_input("你的职位", placeholder="工程师")
                            reviewer_tenure = st.selectbox("在职时长", ["不到1年", "1-2年", "2-3年", "3-5年", "5年以上"])
                        
                        title = st.text_input("评价标题", placeholder="一句话总结")
                        pros = st.text_area("优点", placeholder="公司有什么好的地方？")
                        cons = st.text_area("缺点", placeholder="公司有什么不好的地方？")
                        
                        pitfall_options = ["996", "PUA", "内卷", "欠薪", "裁员", "画饼", "克扣"]
                        selected_pitfalls = st.multiselect("坑点标签（可选）", pitfall_options)
                        
                        if st.form_submit_button("提交评价", type="primary"):
                            if pros and cons:
                                from src.jobgraph.user.contribution import contribution_manager
                                
                                result = contribution_manager.submit_review(
                                    company_id=company_id,
                                    overall_rating=overall_rating,
                                    pros=pros,
                                    cons=cons,
                                    title=title,
                                    salary_rating=salary_rating,
                                    work_life_rating=work_life_rating,
                                    management_rating=management_rating,
                                    reviewer_title=reviewer_title,
                                    reviewer_tenure=reviewer_tenure,
                                    pitfall_tags=selected_pitfalls,
                                )
                                
                                if result["success"]:
                                    st.success(f"✅ 评价已提交！获得 10 积分")
                                else:
                                    st.error(f"提交失败: {result['error']}")
                            else:
                                st.warning("请填写优点和缺点")
    
    with tab2:
        st.subheader("提交坑点")
        
        company_query_pitfall = st.text_input("搜索公司", placeholder="输入公司名称", key="pitfall_company")
        
        if company_query_pitfall:
            companies = job_manager.search_companies(company_query_pitfall)
            if companies:
                company_options = [c.get("c", {}).get("name", "") for c in companies]
                selected_company = st.selectbox("选择公司", company_options, key="pitfall_select")
                
                company_id = None
                for c in companies:
                    if c.get("c", {}).get("name") == selected_company:
                        company_id = c.get("c", {}).get("id")
                        break
                
                if company_id:
                    with st.form("pitfall_form"):
                        pitfall_type = st.selectbox("坑点类型", ["欠薪", "PUA", "996", "内卷", "裁员", "画饼", "克扣", "其他"])
                        severity = st.slider("严重程度", 1, 5, 3)
                        description = st.text_area("详细描述", placeholder="请详细描述坑点情况")
                        evidence = st.text_area("证据来源（可选）", placeholder="如：脉脉评价、劳动仲裁等")
                        
                        if st.form_submit_button("提交坑点", type="primary"):
                            if description:
                                from src.jobgraph.user.contribution import contribution_manager
                                
                                result = contribution_manager.submit_pitfall(
                                    company_id=company_id,
                                    pitfall_type=pitfall_type,
                                    description=description,
                                    severity=severity,
                                    evidence=evidence,
                                )
                                
                                if result["success"]:
                                    st.success(f"✅ 坑点已提交！获得 15 积分")
                                else:
                                    st.error(f"提交失败: {result['error']}")
                            else:
                                st.warning("请填写详细描述")
    
    with tab3:
        st.subheader("提交薪资信息")
        
        company_query_salary = st.text_input("搜索公司", placeholder="输入公司名称", key="salary_company")
        
        if company_query_salary:
            companies = job_manager.search_companies(company_query_salary)
            if companies:
                company_options = [c.get("c", {}).get("name", "") for c in companies]
                selected_company = st.selectbox("选择公司", company_options, key="salary_select")
                
                company_id = None
                for c in companies:
                    if c.get("c", {}).get("name") == selected_company:
                        company_id = c.get("c", {}).get("id")
                        break
                
                if company_id:
                    with st.form("salary_form"):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            job_title = st.text_input("职位名称", placeholder="后端工程师")
                            salary_min = st.number_input("最低薪资 (K)", 0, 200, 20)
                        
                        with col2:
                            experience_years = st.number_input("工作年限", 0, 30, 3)
                            salary_max = st.number_input("最高薪资 (K)", 0, 200, 40)
                        
                        education = st.selectbox("学历", ["大专", "本科", "硕士", "博士"])
                        
                        if st.form_submit_button("提交薪资", type="primary"):
                            if job_title:
                                from src.jobgraph.user.contribution import contribution_manager
                                
                                result = contribution_manager.submit_salary(
                                    company_id=company_id,
                                    job_title=job_title,
                                    salary_min=salary_min * 1000,
                                    salary_max=salary_max * 1000,
                                    experience_years=experience_years,
                                    education=education,
                                )
                                
                                if result["success"]:
                                    st.success(f"✅ 薪资信息已提交！获得 5 积分")
                                else:
                                    st.error(f"提交失败: {result['error']}")
                            else:
                                st.warning("请填写职位名称")
    
    with tab4:
        st.subheader("提交职位信息")
        
        company_query_job = st.text_input("搜索公司", placeholder="输入公司名称", key="job_company")
        
        if company_query_job:
            companies = job_manager.search_companies(company_query_job)
            if companies:
                company_options = [c.get("c", {}).get("name", "") for c in companies]
                selected_company = st.selectbox("选择公司", company_options, key="job_select")
                
                company_id = None
                for c in companies:
                    if c.get("c", {}).get("name") == selected_company:
                        company_id = c.get("c", {}).get("id")
                        break
                
                if company_id:
                    with st.form("job_form"):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            job_title = st.text_input("职位名称 *", placeholder="后端工程师")
                            location = st.text_input("工作地点", placeholder="北京")
                            salary_min = st.number_input("最低薪资 (K)", 0, 200, 20)
                            salary_max = st.number_input("最高薪资 (K)", 0, 200, 40)
                        
                        with col2:
                            experience_years = st.number_input("工作年限要求", 0, 30, 3)
                            education = st.selectbox("学历要求", ["不限", "大专", "本科", "硕士", "博士"])
                            skills = st.text_input("技能要求 (逗号分隔)", placeholder="Java, Python, MySQL")
                            benefits = st.text_input("福利待遇 (逗号分隔)", placeholder="五险一金, 年终奖")
                        
                        description = st.text_area("职位描述", placeholder="请描述职位职责和要求...")
                        
                        if st.form_submit_button("提交职位", type="primary"):
                            if job_title:
                                from src.jobgraph.user.contribution import contribution_manager
                                
                                result = contribution_manager.submit_job(
                                    company_id=company_id,
                                    company_name=selected_company,
                                    title=job_title,
                                    location=location,
                                    salary_min=salary_min * 1000 if salary_min > 0 else None,
                                    salary_max=salary_max * 1000 if salary_max > 0 else None,
                                    experience_years=experience_years,
                                    education=education if education != "不限" else None,
                                    skills=[s.strip() for s in skills.split(",") if s.strip()],
                                    description=description,
                                    benefits=[b.strip() for b in benefits.split(",") if b.strip()],
                                )
                                
                                if result["success"]:
                                    st.success(f"✅ 职位已提交！获得 8 积分")
                                else:
                                    st.error(f"提交失败: {result['error']}")
                            else:
                                st.warning("请填写职位名称")


# ============================================================
# User Center Page
# ============================================================

elif page == "👤 用户中心":
    st.header("👤 用户中心")
    
    user_stats = user_manager.get_user_stats()
    
    # 用户信息卡片
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("基本信息")
        st.write(f"**昵称**: {user_stats['nickname']}")
        st.write(f"**设备ID**: {user_stats['device_id']}")
        st.write(f"**等级**: Lv.{user_stats['level']}")
        st.write(f"**积分**: {user_stats['points']}")
    
    with col2:
        st.subheader("贡献统计")
        contributions = user_stats.get("contributions", {})
        st.write(f"**评价**: {contributions.get('reviews', 0)} 条")
        st.write(f"**坑点**: {contributions.get('pitfalls', 0)} 条")
        st.write(f"**薪资**: {contributions.get('salaries', 0)} 条")
    
    st.divider()
    
    # 设置
    st.subheader("设置")
    
    new_nickname = st.text_input("修改昵称", value=user_stats['nickname'])
    if st.button("保存昵称"):
        user_manager.set_nickname(new_nickname)
        st.success("昵称已更新")
        st.rerun()


# ============================================================
# Data Sync Page
# ============================================================

elif page == "🔄 数据同步":
    st.header("🔄 数据同步")
    
    st.info("从数据中心同步最新的公司、岗位、评价数据")
    
    # 导入自动同步模块
    from src.jobgraph.sync.auto_sync import auto_sync
    
    # 获取同步状态
    sync_status = auto_sync.get_status_info()
    
    # 显示同步状态
    st.subheader("📊 同步状态")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        connected = sync_status.get("connected", False)
        st.metric("连接状态", "✅ 已连接" if connected else "❌ 未连接")
    with col2:
        st.metric("同步次数", sync_status.get("sync_count", 0))
    with col3:
        st.metric("公司数量", sync_status.get("companies_synced", 0))
    with col4:
        st.metric("职位数量", sync_status.get("jobs_synced", 0))
    
    last_sync = sync_status.get("last_sync")
    if last_sync:
        st.caption(f"上次同步: {last_sync[:19]}")
    
    st.divider()
    
    # 同步模式选择
    sync_mode = st.radio(
        "选择同步模式",
        ["🔄 自动同步", "📦 离线数据包", "🌐 手动同步"],
        horizontal=True,
    )
    
    st.divider()
    
    # 自动同步
    if sync_mode == "🔄 自动同步":
        st.subheader("🔄 自动同步配置")
        
        with st.form("auto_sync_config"):
            server_url = st.text_input(
                "数据中心地址",
                value=auto_sync.config.get("server_url", ""),
                placeholder="http://192.168.x.x:8000",
                help="admin 仓库的 API 服务地址"
            )
            
            col1, col2 = st.columns(2)
            with col1:
                auto_sync_enabled = st.checkbox("启用自动同步", value=auto_sync.config.get("auto_sync", False))
                sync_on_startup = st.checkbox("启动时同步", value=auto_sync.config.get("sync_on_startup", True))
            with col2:
                sync_interval = st.selectbox(
                    "同步频率",
                    ["每天", "每周", "每月"],
                    index=0,
                )
            
            if st.form_submit_button("💾 保存配置"):
                interval_map = {"每天": 86400, "每周": 604800, "每月": 2592000}
                auto_sync.save_config({
                    "server_url": server_url,
                    "auto_sync": auto_sync_enabled,
                    "sync_on_startup": sync_on_startup,
                    "sync_interval": interval_map.get(sync_interval, 86400),
                })
                auto_sync.server_url = server_url
                st.success("✅ 配置已保存")
                st.rerun()
        
        st.divider()
        
        # 立即同步按钮
        if st.button("🔄 立即同步", type="primary"):
            if not auto_sync.server_url:
                st.warning("请先配置数据中心地址")
            else:
                with st.spinner("正在同步..."):
                    result = auto_sync.check_and_sync()
                    
                    if result.get("success"):
                        st.success(f"✅ 同步完成！")
                        st.write(f"- 公司: {result.get('companies', 0)} 家")
                        st.write(f"- 职位: {result.get('jobs', 0)} 个")
                        st.write(f"- 评价: {result.get('reviews', 0)} 条")
                        st.rerun()
                    else:
                        st.error(f"同步失败: {result.get('error', '未知错误')}")
    
    # 离线数据包
    elif sync_mode == "📦 离线数据包":
        st.subheader("📦 离线数据包导入")
        
        st.markdown("""
        **使用场景**: 
        - 从数据管理员处获取数据包
        
        **操作步骤**:
        1. 从数据管理员获取 `.json` 数据包
        2. 上传数据包
        3. 点击导入
        """)
        
        uploaded_file = st.file_uploader("上传数据包", type=["json"])
        
        if uploaded_file and st.button("导入数据包", type="primary"):
            with st.spinner("正在导入..."):
                try:
                    import tempfile
                    import os
                    
                    # 保存上传文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as tmp:
                        tmp.write(uploaded_file.getbuffer())
                        tmp_path = tmp.name
                    
                    # 导入
                    import json
                    with open(tmp_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    
                    # 导入到 Neo4j
                    from src.graph.neo4j_client import neo4j_client
                    
                    companies_count = 0
                    jobs_count = 0
                    reviews_count = 0
                    
                    # 导入公司
                    for company in data.get("companies", []):
                        try:
                            cypher = """
                            MERGE (c:Company {id: $id})
                            SET c.name = $name, c.industry = $industry, c.updated_at = datetime()
                            """
                            neo4j_client.execute_write(cypher, company)
                            companies_count += 1
                        except Exception as e:
                            logger.error(f"导入公司失败: {e}")
                    
                    # 导入职位
                    for job in data.get("jobs", []):
                        try:
                            cypher = """
                            MERGE (j:Job {id: $id})
                            SET j.title = $title, j.company_name = $company_name,
                                j.description = $description, j.requirements = $requirements,
                                j.updated_at = datetime()
                            WITH j
                            MATCH (c:Company {id: $company_id})
                            MERGE (c)-[:HAS_JOB]->(j)
                            """
                            neo4j_client.execute_write(cypher, job)
                            jobs_count += 1
                        except Exception as e:
                            logger.error(f"导入职位失败: {e}")
                    
                    # 导入评价
                    for review in data.get("reviews", []):
                        try:
                            cypher = """
                            MERGE (r:Review {id: $id})
                            SET r.company_id = $company_id, r.title = $title,
                                r.overall_rating = $overall_rating, r.updated_at = datetime()
                            WITH r
                            MATCH (c:Company {id: $company_id})
                            MERGE (c)-[:HAS_REVIEW]->(r)
                            """
                            neo4j_client.execute_write(cypher, review)
                            reviews_count += 1
                        except Exception as e:
                            logger.error(f"导入评价失败: {e}")
                    
                    # 清理
                    os.unlink(tmp_path)
                    
                    st.success("✅ 数据包导入成功！")
                    st.write(f"- 公司: {companies_count} 家")
                    st.write(f"- 职位: {jobs_count} 个")
                    st.write(f"- 评价: {reviews_count} 条")
                    
                except Exception as e:
                    st.error(f"导入失败: {e}")
    
    # 手动同步
    elif sync_mode == "🌐 手动同步":
        st.subheader("🌐 手动同步")
        
        server_url = st.text_input(
            "数据中心地址",
            value=auto_sync.server_url,
            placeholder="http://192.168.x.x:8000",
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("测试连接"):
                if server_url:
                    auto_sync.server_url = server_url
                    with st.spinner("测试中..."):
                        if auto_sync.test_connection():
                            st.success("✅ 连接成功！")
                        else:
                            st.error("❌ 连接失败")
                else:
                    st.warning("请输入数据中心地址")
        
        with col2:
            if st.button("立即同步", type="primary"):
                if server_url:
                    auto_sync.server_url = server_url
                    with st.spinner("同步中..."):
                        result = auto_sync.check_and_sync()
                        
                        if result.get("success"):
                            st.success("✅ 同步完成！")
                            st.json(result)
                            st.rerun()
                        else:
                            st.error(f"同步失败: {result.get('error', '未知错误')}")
                else:
                    st.warning("请输入数据中心地址")


# ============================================================
# LLM Configuration
# ============================================================

elif page == "⚙️ LLM 配置":
    st.header("⚙️ LLM 配置")
    
    # 当前状态
    is_configured = config_manager.is_llm_configured()
    
    if is_configured:
        st.success("✅ LLM 已配置 - 简历解析将使用 AI 智能提取")
    else:
        st.warning("⚠️ LLM 未配置 - 简历解析使用规则模式（精度有限）")
    
    st.divider()
    
    # 获取当前配置
    llm_config = config_manager.get_llm_config()
    
    # 选择提供商
    provider = st.radio(
        "选择 LLM 提供商",
        ["OpenAI API", "本地 Ollama"],
        horizontal=True,
    )
    
    if provider == "OpenAI API":
        st.subheader("OpenAI API 配置")
        
        st.info("""
        **获取 API Key**：
        1. 访问 https://platform.openai.com/api-keys
        2. 创建新的 API Key
        3. 复制并粘贴到下方
        
        **兼容 API**：
        - 支持 OpenAI 兼容的第三方 API（如 DeepSeek、Moonshot 等）
        - 只需修改 API Base URL 和 Model 名称
        """)
        
        with st.form("openai_config"):
            openai_key = st.text_input(
                "API Key *",
                value=llm_config.get("openai_api_key", ""),
                type="password",
                placeholder="sk-xxxxxxxxxxxxxxxx",
                help="OpenAI API Key 或兼容 API 的 Key"
            )
            
            col1, col2 = st.columns(2)
            
            with col1:
                openai_base = st.text_input(
                    "API Base URL",
                    value=llm_config.get("openai_api_base", "https://api.openai.com/v1"),
                    help="OpenAI 或兼容 API 的地址"
                )
            
            with col2:
                openai_model = st.text_input(
                    "模型名称",
                    value=llm_config.get("openai_model", "gpt-4o"),
                    help="如 gpt-4o, gpt-3.5-turbo, deepseek-chat 等"
                )
            
            if st.form_submit_button("💾 保存配置", type="primary"):
                if openai_key and openai_key != "sk-your-openai-api-key":
                    success = config_manager.save_llm_config(
                        provider="openai",
                        openai_api_key=openai_key,
                        openai_api_base=openai_base,
                        openai_model=openai_model,
                    )
                    if success:
                        st.success("✅ 配置已保存！重启服务后生效")
                        st.rerun()
                    else:
                        st.error("❌ 保存失败")
                else:
                    st.warning("请输入有效的 API Key")
        
        # 常用 API 预设（在表单外面）
        st.write("**快速设置**：")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("OpenAI 官方"):
                config_manager.save_llm_config(
                    provider="openai",
                    openai_api_base="https://api.openai.com/v1",
                    openai_model="gpt-4o",
                )
                st.rerun()
        
        with col2:
            if st.button("DeepSeek"):
                config_manager.save_llm_config(
                    provider="openai",
                    openai_api_base="https://api.deepseek.com/v1",
                    openai_model="deepseek-chat",
                )
                st.rerun()
        
        with col3:
            if st.button("Moonshot"):
                config_manager.save_llm_config(
                    provider="openai",
                    openai_api_base="https://api.moonshot.cn/v1",
                    openai_model="moonshot-v1-8k",
                )
                st.rerun()
    
    else:  # Ollama
        st.subheader("本地 Ollama 配置")
        
        st.info("""
        **安装 Ollama**：
        ```bash
        # Linux/macOS
        curl -fsSL https://ollama.com/install.sh | sh
        
        # 下载模型
        ollama pull qwen2.5:14b
        ```
        
        **推荐模型**：
        - `qwen2.5:14b` - 中文效果好，需要 16GB 内存
        - `qwen2.5:7b` - 轻量版，需要 8GB 内存
        - `llama3:8b` - 英文效果好
        """)
        
        with st.form("ollama_config"):
            col1, col2 = st.columns(2)
            
            with col1:
                ollama_url = st.text_input(
                    "Ollama 服务地址",
                    value=llm_config.get("ollama_base_url", "http://localhost:11434"),
                    help="Ollama 服务的地址"
                )
            
            with col2:
                ollama_model = st.text_input(
                    "模型名称",
                    value=llm_config.get("ollama_model", "qwen2.5:14b"),
                    help="已下载的模型名称"
                )
            
            if st.form_submit_button("💾 保存配置", type="primary"):
                success = config_manager.save_llm_config(
                    provider="ollama",
                    ollama_base_url=ollama_url,
                    ollama_model=ollama_model,
                )
                if success:
                    st.success("✅ 配置已保存！重启服务后生效")
                    st.rerun()
                else:
                    st.error("❌ 保存失败")
    
    st.divider()
    
    # 测试连接
    st.subheader("🧪 测试连接")
    
    # 显示当前激活的提供商
    if config_manager.is_llm_configured():
        active_provider = config_manager.get_active_provider()
        st.info(f"当前激活的 LLM 提供商: **{active_provider.upper()}**")
    
    if st.button("测试 LLM 连接"):
        if not config_manager.is_llm_configured():
            st.warning("请先配置 LLM（输入 OpenAI API Key 或配置 Ollama 地址）")
        else:
            with st.spinner("正在测试连接..."):
                try:
                    from src.jobgraph.resume.extractor import resume_extractor
                    
                    # 重置 LLM 可用性缓存
                    resume_extractor._llm_available = None
                    
                    # 测试提取
                    test_text = "5年Java开发经验，熟悉Spring Boot、MySQL、Redis"
                    result = resume_extractor.extract(test_text)
                    
                    st.success("✅ LLM 连接成功！")
                    st.write(f"测试结果：识别到 {len(result.skills)} 个技能")
                    st.write(f"技能: {result.skills}")
                except Exception as e:
                    st.error(f"❌ 连接失败: {e}")
                    st.info("请检查：\n- API Key 是否正确\n- Ollama 服务是否启动\n- 网络连接是否正常")
    
    st.divider()
    
    # 隐私说明
    st.info("""
    🔒 **隐私说明**：
    - API Key 仅保存在本地 `.env` 文件中
    - 简历内容通过 API 发送到 LLM 服务进行解析
    - 如果担心隐私，建议使用本地 Ollama
    """)


# ============================================================
# Footer
# ============================================================

st.divider()
st.caption("🔒 JobGraph - 完全免费 | 你的求职，你做主 | 数据仅存在本地，不留任何痕迹")
