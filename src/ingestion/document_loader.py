"""Document loader for multiple file formats."""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from loguru import logger


@dataclass
class Document:
    """Loaded document representation."""

    id: str
    title: str
    source_path: str
    file_type: str
    content: str
    metadata: dict = field(default_factory=dict)
    page_count: Optional[int] = None


class DocumentLoader:
    """Load documents from various file formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}

    def load(self, file_path: str | Path) -> Document:
        """Load a document from file path."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}")

        logger.info(f"Loading document: {path.name}")

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext in (".docx", ".doc"):
            return self._load_docx(path)
        elif ext in (".md", ".markdown"):
            return self._load_markdown(path)
        elif ext == ".txt":
            return self._load_text(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _load_pdf(self, path: Path) -> Document:
        """Load PDF document."""
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            pages.append(page.get_text())

        content = "\n\n".join(pages)
        doc.close()

        return Document(
            id=self._generate_id(path),
            title=path.stem,
            source_path=str(path),
            file_type="pdf",
            content=content,
            page_count=len(pages),
            metadata={"pages": pages},
        )

    def _load_docx(self, path: Path) -> Document:
        """Load DOCX document."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return Document(
            id=self._generate_id(path),
            title=path.stem,
            source_path=str(path),
            file_type="docx",
            content=content,
            metadata={"paragraph_count": len(paragraphs)},
        )

    def _load_markdown(self, path: Path) -> Document:
        """Load Markdown document."""
        content = path.read_text(encoding="utf-8")

        return Document(
            id=self._generate_id(path),
            title=path.stem,
            source_path=str(path),
            file_type="markdown",
            content=content,
        )

    def _load_text(self, path: Path) -> Document:
        """Load plain text document."""
        content = path.read_text(encoding="utf-8")

        return Document(
            id=self._generate_id(path),
            title=path.stem,
            source_path=str(path),
            file_type="txt",
            content=content,
        )

    def _generate_id(self, path: Path) -> str:
        """Generate unique document ID."""
        import hashlib

        return hashlib.md5(str(path).encode()).hexdigest()[:16]
